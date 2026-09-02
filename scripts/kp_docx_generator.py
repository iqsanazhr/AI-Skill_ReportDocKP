#!/usr/bin/env python3
"""
KP DOCX GENERATOR - Universitas Jenderal Soedirman
Converts Markdown Practical Work (KP) reports into publication-grade Microsoft Word (.docx) documents
that strictly adhere to academic thesis formatting standards of the Faculty of Engineering.
"""

import os
import re
import sys
import argparse
import base64
import urllib.request
import docx
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_table_borders(table, color="B0B0B0", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideV w:val="none"/>
                <w:left w:val="none"/>
                <w:right w:val="none"/>
            </w:tblBorders>
        ''')
        tblPr[0].append(borders)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def is_table_separator(row_str):
    cells = [c.strip() for c in row_str.split('|')[1:-1]]
    if not cells:
        return False
    for c in cells:
        if not c or not all(ch in '-: ' for ch in c):
            return False
    return True

def setup_page_number_footer(section, fmt="decimal", start_num=1, align=WD_ALIGN_PARAGRAPH.RIGHT):
    ns = nsdecls('w')
    section.footer.is_linked_to_previous = False
    pg_xml = parse_xml(f'<w:pgNumType {ns} w:fmt="{fmt}" w:start="{start_num}"/>')
    section._sectPr.append(pg_xml)
    
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    
    r = p.add_run()
    r._r.append(parse_xml(f'<w:fldChar {ns} w:fldCharType="begin"/>'))
    r._r.append(parse_xml(f'<w:instrText {ns} xml:space="preserve"> PAGE </w:instrText>'))
    r._r.append(parse_xml(f'<w:fldChar {ns} w:fldCharType="separate"/>'))
    r._r.append(parse_xml(f'<w:fldChar {ns} w:fldCharType="end"/>'))
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

def add_toc_line(doc, title, page_str, level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.tab_stops.add_tab_stop(Cm(14.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    
    if level == 0:
        p.paragraph_format.left_indent = Cm(0)
    elif level == 1:
        p.paragraph_format.left_indent = Cm(0.5)
    elif level == 2:
        p.paragraph_format.left_indent = Cm(1.0)
    elif level == 3:
        p.paragraph_format.left_indent = Cm(1.5)
        
    run_t = p.add_run(title)
    run_t.font.name = "Times New Roman"
    run_t.font.size = Pt(11)
    if level == 0 or title.startswith('BAB ') or title.startswith('LAMPIRAN') or title in ['DAFTAR ISI', 'DAFTAR GAMBAR', 'DAFTAR TABEL', 'DAFTAR PUSTAKA', 'PERNYATAAN', 'LEMBAR PENGESAHAN']:
        run_t.bold = True
    
    p.add_run('\t')
    
    run_p = p.add_run(page_str)
    run_p.font.name = "Times New Roman"
    run_p.font.size = Pt(11)
    if level == 0:
        run_p.bold = True
    return p

def add_styled_paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, 
                         space_before=0, space_after=6, line_spacing=1.5, bold=False, italic=False, 
                         font_size=12, font_name="Times New Roman", font_color=RGBColor(0,0,0),
                         first_line_indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    
    if text:
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
        for part in parts:
            if not part:
                continue
            run = p.add_run()
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.color.rgb = font_color
            
            if part.startswith('**') and part.endswith('**'):
                run.text = part[2:-2]
                run.bold = True
                run.italic = italic
            elif part.startswith('*') and part.endswith('*'):
                run.text = part[1:-1]
                run.bold = bold
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run.text = part[1:-1]
                run.font.name = "Consolas"
                run.font.size = Pt(font_size - 1)
                run.font.color.rgb = RGBColor(40, 40, 120)
            else:
                run.text = part
                run.bold = bold
                run.italic = italic
    return p

def render_mermaid_to_png(mermaid_code, output_path):
    """Renders mermaid code to PNG via mermaid.ink service"""
    try:
        b64 = base64.b64encode(mermaid_code.strip().encode('utf-8')).decode('ascii')
        url = f'https://mermaid.ink/img/{b64}'
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=12) as response:
            data = response.read()
            with open(output_path, 'wb') as f:
                f.write(data)
            return True
    except Exception as e:
        print(f"Warning: Could not render mermaid via online API ({e})")
        return False

def generate_docx(input_md, output_docx, author=None, nim=None, title=None, logo_path=None, diagrams_dir=None):
    if not os.path.exists(input_md):
        raise FileNotFoundError(f"Input file not found: {input_md}")
        
    doc = docx.Document()
    
    # ----------------------------------------------------
    # SECTION 1: COVER (NO PAGE NUMBER)
    # ----------------------------------------------------
    sec_cover = doc.sections[0]
    sec_cover.page_width = Cm(21.0)
    sec_cover.page_height = Cm(29.7)
    sec_cover.top_margin = Cm(3.0)
    sec_cover.bottom_margin = Cm(3.0)
    sec_cover.left_margin = Cm(4.0)
    sec_cover.right_margin = Cm(3.0)
    sec_cover.header_distance = Cm(1.5)
    sec_cover.footer_distance = Cm(1.5)

    with open(input_md, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Extract Title, Author, NIM if not provided
    doc_title = title
    doc_author = author
    doc_nim = nim
    
    for idx, l in enumerate(lines[:30]):
        if not doc_title and l.startswith('# '):
            doc_title = l[2:].strip()
        if not doc_author and l.startswith('* **Nama**:'):
            doc_author = l.replace('* **Nama**:', '').strip()
        if not doc_nim and l.startswith('* **NIM**:'):
            doc_nim = l.replace('* **NIM**:', '').strip()

    if not doc_title: doc_title = "LAPORAN KERJA PRAKTIK"
    if not doc_author: doc_author = "MAHASISWA INFORMATIKA"
    if not doc_nim: doc_nim = "H1D024XXX"

    # Default Logo
    script_dir = os.path.dirname(os.path.abspath(__file__))
    if not logo_path or not os.path.exists(logo_path):
        candidate_logo = os.path.join(script_dir, '..', 'assets', 'logo_unsoed.png')
        if os.path.exists(candidate_logo):
            logo_path = candidate_logo

    # Default Diagrams Dir
    if not diagrams_dir:
        diagrams_dir = os.path.join(os.path.dirname(input_md), 'extracted_assets', 'rendered_diagrams')

    diagram_mapping = {
        'Gambar 1.': os.path.join(diagrams_dir, 'diag_arsitektur.png'),
        'Gambar 2.': os.path.join(diagrams_dir, 'diag_use_case.png'),
        'Gambar 3.': os.path.join(diagrams_dir, 'diag_class.png'),
        'Gambar 4.': os.path.join(diagrams_dir, 'diag_erd.png'),
        'Gambar 7.': os.path.join(diagrams_dir, 'diag_flowchart_checkin.png'),
        'Gambar 8.': os.path.join(diagrams_dir, 'diag_seq_checkin.png'),
        'Gambar 13.': os.path.join(diagrams_dir, 'diag_flowchart_konsul.png'),
        'Gambar 14.': os.path.join(diagrams_dir, 'diag_seq_chat.png'),
        'Gambar 19.': os.path.join(diagrams_dir, 'diag_flowchart_api_fcm.png'),
        'Gambar 20.': os.path.join(diagrams_dir, 'diag_seq_fcm.png'),
    }

    # Build Cover
    add_styled_paragraph(
        doc,
        text=doc_title,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=12, space_after=18, line_spacing=1.15,
        bold=True, font_size=14
    )
    
    add_styled_paragraph(
        doc,
        text="LAPORAN KERJA PRAKTIK",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=12, space_after=24, line_spacing=1.15,
        bold=True, font_size=14
    )
    
    if logo_path and os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(12)
        p_logo.paragraph_format.space_after = Pt(24)
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Cm(5.5))
    
    add_styled_paragraph(
        doc,
        text=f"Oleh:\n{doc_author.upper()}\n{doc_nim.upper()}",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=18, space_after=36, line_spacing=1.3,
        bold=True, font_size=12
    )
    
    add_styled_paragraph(
        doc,
        text="KEMENTERIAN PENDIDIKAN TINGGI, SAINS, DAN TEKNOLOGI\nUNIVERSITAS JENDERAL SOEDIRMAN\nFAKULTAS TEKNIK\nJURUSAN INFORMATIKA\nPURBALINGGA\n2026",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=24, space_after=0, line_spacing=1.2,
        bold=True, font_size=14
    )

    # ----------------------------------------------------
    # SECTION 2: FRONT MATTER (ROMAN NUMERALS: i, ii, iii...)
    # ----------------------------------------------------
    sec_front = doc.add_section()
    sec_front.page_width = Cm(21.0)
    sec_front.page_height = Cm(29.7)
    sec_front.top_margin = Cm(3.0)
    sec_front.bottom_margin = Cm(3.0)
    sec_front.left_margin = Cm(4.0)
    sec_front.right_margin = Cm(3.0)
    setup_page_number_footer(sec_front, fmt="lowerRoman", start_num=1, align=WD_ALIGN_PARAGRAPH.RIGHT)

    start_idx = 0
    for idx, l in enumerate(lines):
        if l.strip().startswith('## PERNYATAAN'):
            start_idx = idx
            break

    in_code_block = False
    code_buffer = []
    table_buffer = []
    in_daftar_isi = False
    in_daftar_gambar = False
    in_daftar_tabel = False
    has_switched_to_body_section = False

    def process_table_buffer():
        nonlocal table_buffer
        if not table_buffer:
            return
            
        headers = [c.strip() for c in table_buffer[0].split('|')[1:-1]]
        data_rows = []
        for r in table_buffer[1:]:
            if is_table_separator(r):
                continue
            cols = [c.strip() for c in r.split('|')[1:-1]]
            if all(not c or all(ch in '-: ' for ch in c) for c in cols):
                continue
            data_rows.append(cols)

        if headers and data_rows:
            num_cols = max(len(headers), max(len(row) for row in data_rows))
            is_signature_table = "Dosen Pembimbing" in headers[0] and "Pembimbing Lapangan" in headers[-1]

            tbl = doc.add_table(rows=len(data_rows) + 1, cols=num_cols)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            if not is_signature_table:
                set_table_borders(tbl, color="B0B0B0", sz="4", val="single")

            # Header Row
            hdr_row = tbl.rows[0]
            trPr = hdr_row._tr.get_or_add_trPr()
            trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
            
            for col_idx in range(num_cols):
                cell = hdr_row.cells[col_idx]
                if not is_signature_table:
                    set_cell_background(cell, "F1F5F9")
                set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.0
                text_val = headers[col_idx] if col_idx < len(headers) else ""
                run = p.add_run(text_val)
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(15, 23, 42)

            # Data Rows
            for row_idx, r_data in enumerate(data_rows):
                row = tbl.rows[row_idx + 1]
                trPr_row = row._tr.get_or_add_trPr()
                trPr_row.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
                
                bg_color = "F8FAFC" if (row_idx % 2 == 1 and not is_signature_table) else "FFFFFF"
                for col_idx in range(num_cols):
                    cell = row.cells[col_idx]
                    if bg_color != "FFFFFF":
                        set_cell_background(cell, bg_color)
                    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.15
                    
                    raw_val = r_data[col_idx] if col_idx < len(r_data) else ""
                    lines_in_cell = raw_val.split('<br>')
                    
                    for line_idx, cell_line in enumerate(lines_in_cell):
                        if line_idx > 0:
                            p = cell.add_paragraph()
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(2)
                            p.paragraph_format.line_spacing = 1.15
                        
                        if is_signature_table:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif len(cell_line.strip()) <= 4 or cell_line.strip().isdigit() or cell_line.strip() in ['✓', 'A', '100%']:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            
                        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', cell_line)
                        for part in parts:
                            if not part: continue
                            run = p.add_run()
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(9.5 if not is_signature_table else 11)
                            if part.startswith('**') and part.endswith('**'):
                                run.text = part[2:-2]
                                run.bold = True
                            elif part.startswith('*') and part.endswith('*'):
                                run.text = part[1:-1]
                                run.italic = True
                            elif part.startswith('`') and part.endswith('`'):
                                run.text = part[1:-1]
                                run.font.name = "Consolas"
                                run.font.size = Pt(8.5)
                            else:
                                run.text = part

            p_after = doc.add_paragraph()
            p_after.paragraph_format.space_before = Pt(0)
            p_after.paragraph_format.space_after = Pt(6)

        table_buffer = []

    i = start_idx
    while i < len(lines):
        line = lines[i].rstrip('\r\n')
        stripped = line.strip()

        # Switch to Body Section at BAB I
        if stripped.startswith('## BAB I') and not has_switched_to_body_section:
            process_table_buffer()
            sec_body = doc.add_section()
            sec_body.page_width = Cm(21.0)
            sec_body.page_height = Cm(29.7)
            sec_body.top_margin = Cm(3.0)
            sec_body.bottom_margin = Cm(3.0)
            sec_body.left_margin = Cm(4.0)
            sec_body.right_margin = Cm(3.0)
            setup_page_number_footer(sec_body, fmt="decimal", start_num=1, align=WD_ALIGN_PARAGRAPH.RIGHT)
            has_switched_to_body_section = True

        # Handle Code Block
        if stripped.startswith('```'):
            process_table_buffer()
            if in_code_block:
                in_code_block = False
                code_text = "\n".join(code_buffer)
                is_mermaid = any(k in code_text for k in ['graph TD', 'graph LR', 'classDiagram', 'erDiagram', 'sequenceDiagram', 'flowchart TD', 'flowchart LR'])
                
                if is_mermaid:
                    os.makedirs(diagrams_dir, exist_ok=True)
                    # Peek ahead to find the caption line
                    caption_text = None
                    caption_line_offset = 0
                    for lookahead in range(1, 4):
                        if i + lookahead < len(lines):
                            nxt = lines[i + lookahead].strip()
                            if nxt.startswith('*Gambar ') or nxt.startswith('Gambar '):
                                caption_text = nxt.strip('*_')
                                caption_line_offset = lookahead
                                break
                    
                    # Generate filename from caption or hash
                    safe_name = re.sub(r'[^a-zA-Z0-9]', '_', caption_text[:30]) if caption_text else f"mermaid_{abs(hash(code_text)) % 100000}"
                    img_path = os.path.join(diagrams_dir, f"{safe_name}.png")
                    
                    # Render if not already exists
                    if not os.path.exists(img_path):
                        print(f"Auto-rendering Mermaid diagram: {safe_name}...")
                        render_mermaid_to_png(code_text, img_path)
                        
                    if os.path.exists(img_path):
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_img.paragraph_format.space_before = Pt(12)
                        p_img.paragraph_format.space_after = Pt(6)
                        run_img = p_img.add_run()
                        run_img.add_picture(img_path, width=Cm(13.8))
                        
                        if caption_text:
                            add_styled_paragraph(
                                doc,
                                text=caption_text,
                                align=WD_ALIGN_PARAGRAPH.CENTER,
                                space_before=4, space_after=14, line_spacing=1.15,
                                bold=True, italic=True, font_size=10.5, font_color=RGBColor(30, 41, 59)
                            )
                            # Advance line pointer past the caption so it's not printed twice
                            i += caption_line_offset
                else:
                    p_code = doc.add_paragraph()
                    p_code.paragraph_format.space_before = Pt(6)
                    p_code.paragraph_format.space_after = Pt(6)
                    p_code.paragraph_format.left_indent = Cm(0.5)
                    p_code.paragraph_format.right_indent = Cm(0.5)
                    p_code.paragraph_format.line_spacing = 1.0
                    run = p_code.add_run(code_text)
                    run.font.name = "Consolas"
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(30, 41, 59)
                code_buffer = []
            else:
                in_code_block = True
                code_buffer = []
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Check section states
        if stripped == '## DAFTAR ISI':
            process_table_buffer()
            in_daftar_isi = True
            in_daftar_gambar = False
            in_daftar_tabel = False
            add_styled_paragraph(doc, text="DAFTAR ISI", align=WD_ALIGN_PARAGRAPH.CENTER,
                                 space_before=18, space_after=18, line_spacing=1.15, bold=True, font_size=14)
            i += 1
            continue
        elif stripped == '## DAFTAR GAMBAR':
            process_table_buffer()
            in_daftar_isi = False
            in_daftar_gambar = True
            in_daftar_tabel = False
            add_styled_paragraph(doc, text="DAFTAR GAMBAR", align=WD_ALIGN_PARAGRAPH.CENTER,
                                 space_before=18, space_after=18, line_spacing=1.15, bold=True, font_size=14)
            i += 1
            continue
        elif stripped == '## DAFTAR TABEL':
            process_table_buffer()
            in_daftar_isi = False
            in_daftar_gambar = False
            in_daftar_tabel = True
            add_styled_paragraph(doc, text="DAFTAR TABEL", align=WD_ALIGN_PARAGRAPH.CENTER,
                                 space_before=18, space_after=18, line_spacing=1.15, bold=True, font_size=14)
            i += 1
            continue
        elif stripped.startswith('## BAB I') or stripped == '---':
            if in_daftar_isi or in_daftar_gambar or in_daftar_tabel:
                in_daftar_isi = False
                in_daftar_gambar = False
                in_daftar_tabel = False

        # Format DAFTAR ISI items
        if in_daftar_isi and (stripped.startswith('* ') or stripped.startswith('  * ') or stripped.startswith('    * ')):
            raw = stripped.lstrip(' *-\t')
            m = re.match(r'^(.*?)\s+[—\.\-]+\s+([ivxlcdm\d]+)$', raw, re.IGNORECASE)
            if m:
                title_item = m.group(1).replace('**', '').strip()
                page_item = m.group(2).strip()
                level = 0
                if stripped.startswith('    * '): level = 2
                elif stripped.startswith('  * '): level = 1
                add_toc_line(doc, title_item, page_item, level=level)
                i += 1
                continue

        # Format DAFTAR GAMBAR & TABEL items
        if (in_daftar_gambar or in_daftar_tabel) and (stripped.startswith('* ') or stripped.startswith('- ')):
            raw = stripped.lstrip(' *-\t')
            m = re.match(r'^(.*?)\s+[—\.\-]+\s+([ivxlcdm\d]+)$', raw, re.IGNORECASE)
            if m:
                title_item = m.group(1).replace('**', '').strip()
                page_item = m.group(2).strip()
                add_toc_line(doc, title_item, page_item, level=0)
                i += 1
                continue

        # Handle Markdown Tables
        if stripped.startswith('|') and stripped.endswith('|'):
            table_buffer.append(stripped)
            i += 1
            continue
        elif table_buffer:
            process_table_buffer()

        if not stripped:
            i += 1
            continue

        if stripped == '---':
            process_table_buffer()
            doc.add_page_break()
            i += 1
            continue

        # Headings
        if stripped.startswith('## '):
            process_table_buffer()
            heading_text = stripped[3:].strip()
            if heading_text.startswith('BAB ') or heading_text in ['PERNYATAAN', 'LEMBAR PENGESAHAN', 'DAFTAR ISI', 'DAFTAR GAMBAR', 'DAFTAR TABEL', 'DAFTAR PUSTAKA'] or heading_text.startswith('LAMPIRAN'):
                add_styled_paragraph(
                    doc,
                    text=heading_text,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before=18, space_after=12, line_spacing=1.15,
                    bold=True, font_size=14
                )
            else:
                add_styled_paragraph(
                    doc,
                    text=heading_text,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before=14, space_after=6, line_spacing=1.15,
                    bold=True, font_size=12
                )
            i += 1
            continue

        if stripped.startswith('### '):
            process_table_buffer()
            heading_text = stripped[4:].strip()
            if heading_text in ['PENDAHULUAN', 'TINJAUAN PUSTAKA', 'PELAKSANAAN KERJA PRAKTIK', 'IMPLEMENTASI', 'PENUTUP']:
                add_styled_paragraph(
                    doc,
                    text=heading_text,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before=0, space_after=18, line_spacing=1.15,
                    bold=True, font_size=14
                )
            elif heading_text.startswith('LAMPIRAN') or heading_text in ['SERTIFIKAT KELULUSAN', 'SURAT PENERIMAAN INSTANSI', 'PENILAIAN PELAKSANAAN KERJA PRAKTIK', 'LEMBAR PRESENSI KERJA PRAKTIK', 'LOGBOOK KERJA PRAKTIK', 'DOKUMENTASI USER ACCEPTANCE TESTING (UAT)', 'DOKUMENTASI KEGIATAN', 'CURRICULUM VITAE']:
                add_styled_paragraph(
                    doc,
                    text=heading_text,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before=0, space_after=14, line_spacing=1.15,
                    bold=True, font_size=12
                )
            else:
                add_styled_paragraph(
                    doc,
                    text=heading_text,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before=12, space_after=6, line_spacing=1.15,
                    bold=True, font_size=12
                )
            i += 1
            continue

        if stripped.startswith('#### '):
            process_table_buffer()
            heading_text = stripped[5:].strip()
            add_styled_paragraph(
                doc,
                text=heading_text,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=10, space_after=4, line_spacing=1.15,
                bold=True, font_size=12
            )
            i += 1
            continue

        # Bullet List Items
        if stripped.startswith('* ') or stripped.startswith('- '):
            process_table_buffer()
            item_text = stripped[2:].strip()
            
            matched_fig = None
            for fig_key, fig_path in diagram_mapping.items():
                if fig_key in item_text and os.path.exists(fig_path):
                    matched_fig = (fig_key, fig_path)
                    break

            if matched_fig:
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(12)
                p_img.paragraph_format.space_after = Pt(6)
                run_img = p_img.add_run()
                run_img.add_picture(matched_fig[1], width=Cm(13.8))
                
                clean_caption = item_text.strip('*_')
                add_styled_paragraph(
                    doc,
                    text=clean_caption,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before=4, space_after=14, line_spacing=1.15,
                    bold=True, italic=True, font_size=10.5, font_color=RGBColor(30, 41, 59)
                )
                i += 1
                continue

            p = add_styled_paragraph(
                doc,
                text=item_text,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_before=2, space_after=4, line_spacing=1.5,
                font_size=12
            )
            p.paragraph_format.left_indent = Cm(0.75)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.runs[0].text = "•  " + p.runs[0].text
            i += 1
            continue

        # Numbered List Items
        num_match = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if num_match:
            process_table_buffer()
            num_str = num_match.group(1)
            item_text = num_match.group(2)
            p = add_styled_paragraph(
                doc,
                text=item_text,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_before=2, space_after=4, line_spacing=1.5,
                font_size=12
            )
            p.paragraph_format.left_indent = Cm(0.75)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.runs[0].text = f"{num_str}.  " + p.runs[0].text
            i += 1
            continue

        # Image Caption or Note
        if stripped.startswith('*Gambar ') or stripped.startswith('Gambar '):
            process_table_buffer()
            clean_caption = stripped.strip('*_')
            matched_fig = None
            for fig_key, fig_path in diagram_mapping.items():
                if fig_key in clean_caption and os.path.exists(fig_path):
                    matched_fig = (fig_key, fig_path)
                    break
            
            if matched_fig:
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(12)
                p_img.paragraph_format.space_after = Pt(6)
                run_img = p_img.add_run()
                run_img.add_picture(matched_fig[1], width=Cm(13.8))

            add_styled_paragraph(
                doc,
                text=clean_caption,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=4, space_after=14, line_spacing=1.15,
                bold=True, italic=True, font_size=10.5, font_color=RGBColor(30, 41, 59)
            )
            i += 1
            continue

        # Table caption
        if stripped.startswith('*Table ') or stripped.startswith('Table '):
            process_table_buffer()
            clean_caption = stripped.strip('*_')
            add_styled_paragraph(
                doc,
                text=clean_caption,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=8, space_after=4, line_spacing=1.15,
                bold=True, font_size=11, font_color=RGBColor(15, 23, 42)
            )
            i += 1
            continue

        # Math formula / equation
        if stripped.startswith('$$') and stripped.endswith('$$'):
            process_table_buffer()
            eq_text = stripped[2:-2].strip()
            eq_text = eq_text.replace(r'\text{', '').replace('}', '').replace(r'\frac{', '(').replace(r'}{', ' / ').replace(r'}', ')').replace(r'\quad', '  ').replace(r'\,', ' ').replace(r'\%', '%').replace(r'{,}', ',')
            add_styled_paragraph(
                doc,
                text=eq_text,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=8, space_after=12, line_spacing=1.2,
                bold=True, font_size=12, font_color=RGBColor(30, 41, 59)
            )
            i += 1
            continue

        # Regular Body Paragraphs
        process_table_buffer()
        add_styled_paragraph(
            doc,
            text=stripped,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_before=0, space_after=6, line_spacing=1.5,
            font_size=12,
            first_line_indent=Cm(1.0)
        )
        i += 1

    process_table_buffer()

    # Save with lock fallback
    try:
        doc.save(output_docx)
        print(f"SUCCESS: Document successfully created at: {output_docx}")
    except PermissionError:
        base, ext = os.path.splitext(output_docx)
        alt_output = f"{base}_New{ext}"
        doc.save(alt_output)
        print(f"WARNING: File {output_docx} is locked in Word. Saved to alternative: {alt_output}")

def main():
    parser = argparse.ArgumentParser(description="Generate Word (.docx) report from markdown for Unsoed KP")
    parser.add_argument("--input", "-i", required=True, help="Path to markdown report file (.md)")
    parser.add_argument("--output", "-o", required=True, help="Path to output Word file (.docx)")
    parser.add_argument("--author", "-a", help="Author name")
    parser.add_argument("--nim", "-n", help="Student NIM")
    parser.add_argument("--title", "-t", help="Report Title")
    parser.add_argument("--logo", "-l", help="Path to logo Unsoed PNG")
    parser.add_argument("--diagrams", "-d", help="Directory containing rendered diagram images")
    
    args = parser.parse_args()
    generate_docx(
        input_md=args.input,
        output_docx=args.output,
        author=args.author,
        nim=args.nim,
        title=args.title,
        logo_path=args.logo,
        diagrams_dir=args.diagrams
    )

if __name__ == '__main__':
    main()
